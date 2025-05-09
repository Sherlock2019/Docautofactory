import streamlit as st
import os
import re
from pathlib import Path

st.set_page_config(page_title="🧠 DocAutoFactory", layout="wide")
st.title("🧠 DocAutoFactory – Dynamic RFP & IT Report Generator")

# ──────────────────────────────────────────────────────────────
# 🔧 SETTINGS
MODULES_DIR = "modules"
TEMPLATE_DIR = "templates"
OUTPUT_DIR = "output"
Path(OUTPUT_DIR).mkdir(exist_ok=True)

# ──────────────────────────────────────────────────────────────
# 📥 Upload Word template
st.sidebar.header("📄 Upload Main Template")
uploaded_template = st.sidebar.file_uploader(
    "Upload .docx template with placeholders like {CUSTOMER_NAME}, {DC_NB}...",
    type=["docx"],
    key="main_template",
)
if uploaded_template:
    template_path = os.path.join(TEMPLATE_DIR, uploaded_template.name)
    with open(template_path, "wb") as f:
        f.write(uploaded_template.read())
    st.sidebar.success(f"Uploaded: {uploaded_template.name}")
else:
    st.warning("Please upload a Word (.docx) template to begin.")

# ──────────────────────────────────────────────────────────────
# 📍 Extract dynamic placeholders from template
def extract_placeholders(docx_path):
    import docx
    doc = docx.Document(docx_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    return sorted(set(re.findall(r"{(.*?)}", text)))

# ──────────────────────────────────────────────────────────────
# 📌 Gather user input for each placeholder
def save_upload(uploaded_file, placeholder):
    save_dir = os.path.join(MODULES_DIR, placeholder)
    Path(save_dir).mkdir(parents=True, exist_ok=True)
    with open(os.path.join(save_dir, uploaded_file.name), "wb") as f:
        f.write(uploaded_file.read())
    st.success(f"Saved {uploaded_file.name} for {placeholder}")

def save_textarea(text, placeholder):
    save_dir = os.path.join(MODULES_DIR, placeholder)
    Path(save_dir).mkdir(parents=True, exist_ok=True)
    with open(os.path.join(save_dir, "main.txt"), "w") as f:
        f.write(text)
    st.success(f"Text saved for {placeholder}")

# ──────────────────────────────────────────────────────────────
# 🧠 Step 2 – Process Placeholder Inputs
if uploaded_template:
    st.subheader("🧩 Step 2: Provide Content for Placeholders")
    placeholder_list = extract_placeholders(template_path)

    for idx, ph in enumerate(placeholder_list):
        st.markdown(f"### {ph}")

        if ph in ["CUSTOMER_NAME", "COMPANY_NAME", "PARTNER_NAME", "CITY_NAME"]:
            value = st.text_input(f"Enter value for {ph}:", key=f"text_{ph}")
            if value:
                save_textarea(value, ph)
        else:
            uploaded = st.file_uploader(
                f"Upload file for {ph}",
                key=f"file_{ph}_{idx}",  # ✅ Ensure unique key!
                type=None
            )
            if uploaded:
                save_upload(uploaded, ph)

# ──────────────────────────────────────────────────────────────
# 🛠️ Step 3 – Generate filled document (basic placeholder replacement)
st.subheader("📤 Step 3: Generate Filled Document")

if st.button("📄 Generate Final Document"):
    from docx import Document

    if not uploaded_template:
        st.error("Please upload a Word template first.")
    else:
        doc = Document(template_path)
        placeholders = extract_placeholders(template_path)

        for p in doc.paragraphs:
            for ph in placeholders:
                module_path = os.path.join(MODULES_DIR, ph, "main.txt")
                if os.path.exists(module_path):
                    with open(module_path, "r") as f:
                        value = f.read().strip()
                    p.text = p.text.replace(f"{{{ph}}}", value)

        output_path = os.path.join(OUTPUT_DIR, "Generated_Document.docx")
        doc.save(output_path)
        with open(output_path, "rb") as f:
            st.download_button(
                label="📥 Download Final Document",
                data=f,
                file_name="Generated_Document.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
